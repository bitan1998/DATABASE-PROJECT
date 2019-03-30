import docx

doc = docx.Document()

doc.add_paragraph("This is some text")
new_paragraph = doc.add_paragraph("I love python programming")
new_paragraph.add_run(". This the continuation")

doc.save("Reciepts/new.docx")